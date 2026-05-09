# 
from pathlib import Path
import json
import sys
import re
import os

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH

# --------------------------
# Paths - Safe handling for Vercel deployment
# --------------------------
BASE_DIR = Path.cwd()
TEMPLATE_PATH = BASE_DIR / "templates" / "resume_template.docx"

# JSON path: Check environment variable first (for /tmp on Vercel), fall back to data/
# This allows Vercel to write to /tmp instead of read-only data/ folder
JSON_PATH_ENV = os.environ.get("RESUME_DATA_PATH")
if JSON_PATH_ENV:
    JSON_PATH = Path(JSON_PATH_ENV)
    print(f"[DEBUG] Using JSON path from env: {JSON_PATH}")
else:
    JSON_PATH = BASE_DIR / "data" / "base_content.json"
    print(f"[DEBUG] Using default JSON path: {JSON_PATH}")

# Use /tmp for output on Vercel, local for development
# Vercel serverless doesn't allow persistent file writes outside /tmp
BASE_FOLDER = Path("/tmp/generated_resumes")
BASE_FOLDER.mkdir(parents=True, exist_ok=True)

# Resume final filename
FINAL_FILENAME = "Sahi_Kolukuluri.docx"

# Temporary debug file
OUTPUT_RAW = Path("Tailored_Resume_RAW.docx")  # saved right after docxtpl render

# Debug logging
print(f"[DEBUG] CWD: {BASE_DIR}")
print(f"[DEBUG] Template path: {TEMPLATE_PATH}")
print(f"[DEBUG] JSON path: {JSON_PATH}")
print(f"[DEBUG] Output dir: {BASE_FOLDER}")

# --------------------------
# Helpers
# --------------------------
LIST_STYLE_CANDIDATES = {
    "List Bullet",
    "List Paragraph",
    "Bullet",
    "List Bullet 2",
    "List Bullet 3",
    "Body Text List",
}

BULLET_PREFIXES = ("•", "-", "–", "—", "*")


def tighten_list_paragraph(p):
    """Tighten line spacing and justify bullet paragraphs."""
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def is_list_style(paragraph):
    """Check if a paragraph uses a list-like style."""
    try:
        return paragraph.style and paragraph.style.name in LIST_STYLE_CANDIDATES
    except Exception:
        return False


def remove_paragraph(paragraph):
    """Remove an empty or unwanted paragraph."""
    p = paragraph._element
    parent = p.getparent()
    parent.remove(p)
    p._p = p._element = None


def iter_all_paragraphs(doc: Document):
    """
    Yield paragraphs from the main document and from tables.
    This helps if any resume content is inside Word tables.
    """
    for p in doc.paragraphs:
        yield p

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def apply_bold_markers_to_paragraph(p):
    """
    Convert **keyword** markers into real bold text in Word.

    Example:
    Input text:  Strong in **Java**, **Spring Boot**, and **Apigee**.
    Output doc: Strong in Java, Spring Boot, and Apigee.  # marked words are bold
    """
    text = p.text

    if "**" not in text:
        return

    # Split text by **bold** markers
    parts = re.split(r"(\*\*.*?\*\*)", text)

    # Keep paragraph style before clearing
    paragraph_style = p.style

    # Clear existing runs
    for run in list(p.runs):
        run._element.getparent().remove(run._element)

    p.style = paragraph_style

    for part in parts:
        if not part:
            continue

        run = p.add_run()

        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run.text = part[2:-2]
            run.bold = True
        else:
            run.text = part


def apply_bold_markers(doc: Document):
    """Apply **bold** conversion across the full document."""
    for p in iter_all_paragraphs(doc):
        apply_bold_markers_to_paragraph(p)


def normalize_bullets(doc: Document):
    """Normalize bullets and ensure consistent list formatting."""
    prev_was_list = False
    to_delete = []

    for p in iter_all_paragraphs(doc):
        txt = p.text.strip()

        # Convert manually typed bullets into Word list style
        if txt.startswith(BULLET_PREFIXES):
            raw = p.text
            new_text = raw.lstrip("".join(BULLET_PREFIXES)).lstrip(" \t")
            p.text = new_text if new_text else ""
            try:
                p.style = doc.styles["List Bullet"]
            except KeyError:
                p.style = doc.styles["List Paragraph"]

            txt = p.text.strip()

        # Tighten and justify list paragraphs
        if is_list_style(p):
            tighten_list_paragraph(p)
            if prev_was_list and txt == "":
                to_delete.append(p)
            prev_was_list = True
        else:
            prev_was_list = False

    # Remove empty bullet paragraphs
    for p in to_delete:
        remove_paragraph(p)


def get_next_subfolder(base: Path) -> Path:
    """Find the next numeric subfolder under base, e.g., 1, 2, 3..."""
    base.mkdir(parents=True, exist_ok=True)

    numbers = []
    for child in base.iterdir():
        if child.is_dir() and re.fullmatch(r"\d+", child.name):
            numbers.append(int(child.name))

    next_num = max(numbers, default=0) + 1
    return base / str(next_num)


def main():
    """Main resume generation logic."""
    print(f"[DEBUG] Checking template existence: {TEMPLATE_PATH.exists()}")
    if not TEMPLATE_PATH.exists():
        print(f"Error: Template not found: {TEMPLATE_PATH.resolve()}")
        sys.exit(1)

    print(f"[DEBUG] Checking JSON existence: {JSON_PATH.exists()}")
    if not JSON_PATH.exists():
        print(f"Error: JSON not found: {JSON_PATH.resolve()}")
        sys.exit(1)

    print(f"[DEBUG] Loading JSON from: {JSON_PATH}")
    with JSON_PATH.open("r", encoding="utf-8") as f:
        ctx = json.load(f)

    # --------------------------
    # Prepare numbered folder (inside /tmp for Vercel compatibility)
    # --------------------------
    print(f"[DEBUG] Creating subfolder in: {BASE_FOLDER}")
    target_folder = get_next_subfolder(BASE_FOLDER)
    target_folder.mkdir(parents=True, exist_ok=True)
    print(f"[DEBUG] Target folder: {target_folder}")

    # Define final paths inside that folder
    output_raw = target_folder / "Tailored_Resume_RAW.docx"
    output_final = target_folder / FINAL_FILENAME
    print(f"[DEBUG] Output paths - Raw: {output_raw}, Final: {output_final}")

    # --------------------------
    # Render raw template
    # --------------------------
    print(f"[DEBUG] Rendering template...")
    doc = DocxTemplate(str(TEMPLATE_PATH))
    doc.render(ctx)
    doc.save(str(output_raw))
    print(f"[DEBUG] Raw document saved to: {output_raw}")

    # --------------------------
    # Post-process final document
    # --------------------------
    print(f"[DEBUG] Post-processing document...")
    d = Document(str(output_raw))

    # First normalize bullets because this may rewrite bullet paragraph text
    normalize_bullets(d)

    # Then convert **keyword** markers into real bold formatting
    apply_bold_markers(d)

    # Final spacing cleanup
    for p in iter_all_paragraphs(d):
        if is_list_style(p):
            tighten_list_paragraph(p)

    d.save(str(output_final))
    print(f"[DEBUG] Final document saved to: {output_final}")

    print(f"✓ Rendered OK: {output_final.resolve()}")
    print(f"[DEBUG] Raw pre-fix version at: {output_raw.resolve()}")


if __name__ == "__main__":
    main()